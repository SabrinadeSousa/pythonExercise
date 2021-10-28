from docx import Document

document = Document()

head = document.add_heading("Grade Currícular")
disc = input("Qual a disciplina?: ")
p = document.add_paragraph("Disciplina: ")
p.add_run(disc).bold = True


carga = input("Qual a carga horária?: ")
p = document.add_paragraph("Com carga horária de: ")
p.add_run(carga).bold = True
p.add_run(' horas.')

tab = document.add_table(rows=1, cols=2)
tab.style="Colorful Grid Accent 6"
cels = tab.rows[0].cells
cels[0].text = 'Número'
cels[0].width = 5
cels[1].text = "Assunto"

for numero in range(1, 5):
   assunto = input("Qual é o assunto número-{}? ".format(numero))
   dados = tab.add_row().cells
   dados[0].text = str(numero)
   dados[0].width = 5
   dados[1].text = assunto


document.save("hello_world.docx")
